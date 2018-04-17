# Python to Word by Uncle Engineer
# www.facebook.com/UncleEngineer
# pip install python-docx

from docx import Document
from docx.shared import Inches
# Create document in python
doc = Document()
doc.add_heading('Uncle Engineer')
# Add text to Paragraph
text = '''
Website: http://uncle-engineer.com/python
Facebook: http://www.facebook.com/UncleEngineer

'''
p = doc.add_paragraph(text)
p.add_run('We can guide you to learn python').bold = True

# Add List Bullet
pstyle = 'List Bullet'
doc.add_paragraph('Basic Python 15 hours', style = pstyle)
doc.add_paragraph('Python for Student', style = pstyle)
# Add Picture
doc.add_picture('logo.png',width=Inches(2))

# Save File
filename = 'Uncle Engineer.docx'
doc.save(filename)
print("%s Saved"%(filename))
